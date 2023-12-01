import re
import os
from docx import Document

# Paths
BASE_PATH = "C:\\Users\\kerem\\Pokemon"
SOURCE_FILE = os.path.join(BASE_PATH, "EditedURls.docx")
POKEMON_DUP = os.path.join(BASE_PATH, "PokemonsOnlyDupV1.docx")
POKEMON_ONLY = os.path.join(BASE_PATH, "PokemonsOnlyV1.docx")
POKEMON_LINKS = os.path.join(BASE_PATH, "PokemonsOnlyLinksV1.docx")
POKEMON_FINAL = os.path.join(BASE_PATH, "PokemonListFinal.docx")

# Extract valid links from the source document
source_doc = Document(SOURCE_FILE)
valid_links = []

pattern = re.compile(r"^/wiki/.+?_\(Pok%C3%A9mon\)$")

for para in source_doc.paragraphs:
    link = para.text.strip()
    if pattern.match(link):
        valid_links.append(link)

# Save the links to the duplicate file
doc_dup = Document()
for link in valid_links:
    doc_dup.add_paragraph(link)
doc_dup.save(POKEMON_DUP)

# Remove duplicates while preserving order
seen = set()
unique_links = [x for x in valid_links if x not in seen and not seen.add(x)]

# Save the unique links to the Pokemon-only file
doc_only = Document()
for link in unique_links:
    doc_only.add_paragraph(link)
doc_only.save(POKEMON_ONLY)

# Add the prefix and save to the links file
prefix = "https://m.bulbapedia.bulbagarden.net"
doc_links = Document()
for link in unique_links:
    doc_links.add_paragraph(prefix + link)
doc_links.save(POKEMON_LINKS)

# Now, append the prefix to each entry of PokemonsOnlyV1.docx and save it as PokemonListFinal.docx
doc_final = Document(POKEMON_ONLY)
for para in doc_final.paragraphs:
    para.text = prefix + para.text
doc_final.save(POKEMON_FINAL)
